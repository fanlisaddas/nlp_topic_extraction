import docx


def get_articles(file_path):
    """
    读取 file_path 路径的文本
    :param file_path: 文本路径
    :return: 返回 list 结构的文本
    """
    doc = docx.Document(file_path)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return fulltext


"""
# 读取abbreviations文件内容
def get_words(filename):
    doc = docx.Document(filename)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return select_words(fulltext)


# 选出abbreviation文件的关键字
def select_words(fulltext):
    keywords = []
    for i in range(len(fulltext)):
        # 将字符串按照\t划分
        flag = fulltext[i].split("\t", 1)
        # 取\t前的字符串即为关键字
        keywords.append(flag[0])
    return keywords
"""
